from django.shortcuts import render,redirect
from django.http import HttpResponseRedirect
from django.core.mail import send_mail
from datetime import datetime
import openpyxl as xl
from .models import Tree, Answer
from .forms import TreeForm, AnswerForm, AnswerStatusForm

import docx
import os
now = datetime.now()
count=0


def tree_list(request):
   tree_list = Tree.objects.all().order_by('number')
   return render(request, 'tree_list.html',{'tree_list': tree_list})

def tree_new(request):
    submitted = False
    if request.method=='POST':
        form = TreeForm(request.POST)
        if form.is_valid():
            form.save()
            return HttpResponseRedirect('/tree_new?submitted=True')
    else:
        form = TreeForm
    if 'submitted' in request.GET:
        submitted = True
    return render(request, 'tree_new.html',{'form': form, 'submitted': submitted})

def tree_ind(request, tree_id):
    tree = Tree.objects.get(pk = tree_id)
    return render(request,'tree_ind.html',{'tree': tree})

def tree_update(request, tree_id):
    tree = Tree.objects.get(pk = tree_id)
    form = TreeForm(request.POST or None, instance = tree)
    if form.is_valid():
        form.save()
        return redirect('tree_list')
    return render(request,'tree_update.html',{'tree':tree,'form':form})

def tree_delete_one(request,tree_id):
    tree = Tree.objects.get(pk = tree_id)
    tree.delete()
    tree_list = Tree.objects.all().order_by('number')
    return render(request,'tree_list.html',{'tree_list':tree_list})

def tree_upload(request):
    wb = xl.load_workbook('Tree.xlsx')
    sheet = wb['Sheet1']
    for row in range(1, sheet.max_row + 1):
        number = sheet.cell(row, 1).value
        question = sheet.cell(row, 2).value
        answer1 = sheet.cell(row, 3).value
        answer2 = sheet.cell(row, 4).value
        answer3 = sheet.cell(row, 5).value
        answer4 = sheet.cell(row, 6).value
        answer5 = sheet.cell(row, 7).value
        result1 = sheet.cell(row, 8).value
        result2 = sheet.cell(row, 9).value
        result3 = sheet.cell(row, 10).value
        result4 = sheet.cell(row, 11).value
        result5 = sheet.cell(row, 11).value
        if Tree.objects.filter(number=number).count() == 0:
            new_item = Tree(number=number,question=question,
                            answer1=answer1,answer2=answer2,answer3=answer3,answer4=answer4,answer5=answer5,
                            result1=result1,result2=result2,result3=result3,result4=result4,result5=result5,
                            temp=True,)
            new_item.save()
    tree_list = Tree.objects.filter(temp=True)
    return render(request, 'tree_upload.html', {'tree_list': tree_list})

def tree_delete(request):
    tree = Tree.objects.all()
    tree.delete()
    return render(request,'tree_list.html',{})

def tree_ask(request):
    if request.method=='POST':
        answer_id = request.POST.__getitem__('answer_id')
        question_no = request.POST.__getitem__('question_no')
        answer = request.POST.__getitem__('answer')
        tree = Tree.objects.filter(number=question_no)[0]
        answer_obj = Answer.objects.filter(pk=answer_id)[0]
        question_text = tree.question
        if answer_obj.name is None and question_no == '1':
            name = request.POST.__getitem__('name')
            answer_obj.name = name
        else:
            name = answer_obj.name

        if answer == "ans1":
            answer_text = question_text + " " + tree.answer1
            result = tree.result1
        if answer == "ans2":
            answer_text = question_text + " " + tree.answer2
            result = tree.result2
        if answer == "ans3":
            answer_text = question_text + " " + tree.answer3
            result = tree.result3
        if answer == "ans4":
            answer_text = question_text + " " + tree.answer4
            result = tree.result4
        if answer == "ans5":
            answer_text = question_text + " " + tree.answer5
            result = tree.result5

        if question_no == '1':
            answer_obj.q1 = answer_text
        if question_no == '2':
            answer_obj.q2 = answer_text
        if question_no == '3':
            answer_obj.q3 = answer_text
        if question_no == '4':
            answer_obj.q4 = answer_text
        if question_no == '5':
            answer_obj.q5 = answer_text
        if question_no == '6':
            answer_obj.q6 = answer_text
        if question_no == '7':
            answer_obj.q7 = answer_text
        if question_no == '8':
            answer_obj.q8 = answer_text
        if question_no == '9':
            answer_obj.q9 = answer_text
        if question_no == '10':
            answer_obj.q10 = answer_text
        answer_obj.save()

        if not result.isnumeric():
            answer_obj.outcome = result
            answer_obj.save()
            path = 'answer_ind/' + answer_id
            return redirect(path)

        question_no = result

    else:
        new_item = Answer()
        new_item.save()
        question_no = 1
        answer_id=new_item.pk
        name = ""

    tree = Tree.objects.filter(number=question_no)[0]

    question = tree.question
    answers = [tree.answer1]
    if tree.answer2 is not None:
        answers.append(tree.answer2)
    if tree.answer3 is not None:
        answers.append(tree.answer3)
    if tree.answer4 is not None:
        answers.append(tree.answer4)
    if tree.answer5 is not None:
        answers.append(tree.answer5)


    return render(request,'tree_ask.html',{'name': name,'question':question,
                                           'answers':answers,'answer_id':answer_id,
                                           'question_no':question_no})

def logic_tree_update(request, answer_id):
    item = Answer.objects.filter(pk = answer_id)[0]
    item.q1 = None
    item.q2 = None
    item.q3 = None
    item.q4 = None
    item.q5 = None
    item.q6 = None
    item.q7 = None
    item.q8 = None
    item.q9 = None
    item.q10 =None
    item.save()
    question_no = 1
    answer_id = item.pk
    tree = Tree.objects.filter(number=question_no)[0]
    name = item.name
    question = tree.question
    answers = [tree.answer1]
    if tree.answer2 is not None:
        answers.append(tree.answer2)
    if tree.answer3 is not None:
        answers.append(tree.answer3)
    if tree.answer4 is not None:
        answers.append(tree.answer4)
    if tree.answer5 is not None:
        answers.append(tree.answer5)

    return render(request,'tree_ask.html',{'name': name,'question':question,
                                           'answers':answers,'answer_id':answer_id,
                                           'question_no':question_no})


def tree_event(request):
    return render(request,'tree_event',{})

def tree_not_event(request):
    return render(request,'tree_not_event',{})

def answer_list(request):
    events_reported = Answer.objects.filter(reported=True).order_by('date_identified')
    events_unreported = Answer.objects.filter(reported=False).exclude(outcome="No Event").order_by('date_identified')
    non_events = Answer.objects.filter(outcome="No Event").order_by('date_identified')
    return render(request, 'answer_list.html',{'events_unreported': events_unreported,
                                               'events_reported': events_reported,
                                               'non_events':non_events})

def answer_new(request):
    submitted = False
    if request.method=='POST':
        form = AnswerForm(request.POST)
        if form.is_valid():
            form.save()
            return HttpResponseRedirect('/answer_new?submitted=True')
    else:
        form = AnswerForm
    if 'submitted' in request.GET:
        submitted = True
    return render(request, 'answer_new.html',{'form': form, 'submitted': submitted})

def answer_ind(request, answer_id):
    answer = Answer.objects.get(pk = answer_id)
    return render(request,'answer_ind.html',{'answer': answer})

def answer_send_emails(request):
    events_reported = Answer.objects.filter(reported=True).order_by('date_identified')
    events_unreported = Answer.objects.filter(reported=False).order_by('date_identified')
    for event in events_unreported:
        print(event)
        print(event.age())
        print(event.email_day())
        if event.age() == event.email_day():
            message = event.next_step()
            send_mail(
                "Query: " + str(event.pk),              #Subject
                message,                                #Message
                "darrenandamanda.robinson@gmail.com",   #Fromemail
                ["darrenandamanda.robinson@gmail.com"], #To email
                fail_silently=False,
            )
    return render(request, 'answer_list.html',{'events_unreported': events_unreported,'events_reported': events_reported})

def answer_send_email(request,answer_id):
    answer = Answer.objects.filter(pk = answer_id)[0]
    send_mail(
        "Query: " + str(answer.pk),               #Subject
        "Please update this query",               #Message
        "darrenandamanda.robinson@gmail.com",     #Fromemail
        ["darrenandamanda.robinson@gmail.com"],   #To email
        fail_silently=False,
    )
    return render(request, 'answer_ind.html',{'answer':answer})

def answer_update(request, answer_id):
    answer = Answer.objects.get(pk = answer_id)
    form = AnswerForm(request.POST or None, instance = answer)
    if form.is_valid():
        form.save()
        return redirect('answer_list')
    return render(request,'answer_update.html',{'answer':answer,'form':form})

def answer_status_update(request, answer_id):
    answer = Answer.objects.get(pk = answer_id)
    form = AnswerStatusForm(request.POST or None, instance = answer)
    if form.is_valid():
        form.save()
        return redirect('answer_list')
    return render(request,'answer_status_update.html',{'answer':answer,'form':form})

def answer_upload(request):
    wb = xl.load_workbook('Answer.xlsx')
    sheet = wb['Sheet1']
    for row in range(1, sheet.max_row + 1):
        number = sheet.cell(row, 1).value
        name = sheet.cell(row, 2).value
        owner = sheet.cell(row, 3).value
        if People.objects.filter(name=owner).count() > 0:
            owner= People.objects.filter(name=owner)[0]
        else:
            owner = People.objects.filter(name='Unallocated')[0]
        if Answer.objects.filter(name=name).count() == 0:
            new_item = Answer(number=number,owner=owner, temp=True)
            new_item.save()
    answer_list = Answer.objects.filter(temp=True)
    return render(request, 'answer_upload.html', {'answer_list': answer_list})

def answer_delete_one(request,answer_id):
    answer = Answer.objects.get(pk = answer_id)
    answer.delete()
    events_reported = Answer.objects.filter(reported=True).order_by('date_identified')
    events_unreported = Answer.objects.filter(reported=False).order_by('date_identified')
    return render(request, 'answer_list.html',{'events_unreported': events_unreported,'events_reported': events_reported})

def answer_delete(request):
    answer = Answer.objects.all()
    answer.delete()
    return render(request,'answer_list.html',{})

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def answer_print(request, answer_id):
    answer = Answer.objects.filter(pk=answer_id)[0]
    doc = docx.Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'

    doc.add_paragraph("Reportable Situation Form")
    doc.add_paragraph("Short name:" + answer.name)

    records = (
        ("Date occurred",str(answer.date_arose)),
        ("Date identified",str(answer.date_identified)),
        ("Nature of the reportable situation", answer.nature),
        ("Description of the reportable situation", answer.description),
        ("Why the breach is significant", answer.why_reportable),
        ("How the reportable situation was identified", answer.how_identified),
        ("How long the duration lasted", answer.duration),
        ("Information about representatives", answer.representative_details),
        ("Whether and how the reportable situation has been rectified", answer.how_rectified),
        ("Whether and when affected clients have been compensated", answer.remediation),
        ("Future compliance", answer.future_compliance),
    )

    table = doc.add_table(rows=1,cols=2)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Item"
    header_cells[1].text = "Description"
    header_cells[0].width = 800000
    header_cells[1].width = 5000000
    for item, description in records:
        row_cells =table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = description
        row_cells[0].width = 800000
        row_cells[1].width = 5000000

    make_rows_bold(table.rows[0])
    doc_name = "Report" + str(answer.pk) + ".docx"
    doc.save(doc_name)
    os.system("start " + doc_name)
    events_reported = Answer.objects.filter(reported=True).order_by('date_identified')
    events_unreported = Answer.objects.filter(reported=False).order_by('date_identified')
    return render(request, 'answer_list.html',{'events_unreported': events_unreported,'events_reported': events_reported})

